#!/usr/bin/env python3
"""docx → markdown 변환기 — 맘곁 태교 실천편용.

8개 docx 입력 → docs/practice/*.md 출력. 빌더가 다른 모듈에서 이 디렉터리를
이론편과 같은 방식으로 빌드해 practice.html / practice-modern.html을 만든다.
"""

from __future__ import annotations
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SRC = Path("/Users/euicheolkwon/Desktop/맘곁태교(실천편)")
OUT = Path(__file__).parent / "docs" / "practice"
OUT.mkdir(parents=True, exist_ok=True)

# (docx 파일명, 출력 마크다운 파일명, H1 제목)
JOBS = [
    ("02_프롤로그.docx",                    "prologue.md",  "들어가며"),
    ("03_책전체구성.docx",                   "00-overview.md","책 전체 구성"),
    ("PART1_태아이전의마음준비_전자책.docx", "part1.md",     "태아 이전의 마음 준비"),
    ("PART2_임신초기태교_전자책.docx",       "part2.md",     "임신 초기 태교"),
    ("PART3_임신중기태교_전자책.docx",       "part3.md",     "임신 중기 태교"),
    ("PART4_임신후기태교_전자책.docx",       "part4.md",     "임신 후기 태교"),
    ("PART5_출산과첫만남_전자책.docx",       "part5.md",     "출산과 첫 만남"),
    ("04_에필로그.docx",                    "epilogue.md",  "닫으며"),
]

# 시각적 섹션 마커 (docx 본문에서 H1/H2 스타일이 안 잡힌 emoji-led headers)
EMOJI_SECTION_MARKERS = {
    "💭": "예비 부모들의 질문",
    "🔬": "현대 과학의 발견",
    "🗣️": "오늘의 태담",
    "✨": "핵심 요약",
    "💡": "오늘의 한 가지 실천",
    "📋": "실천 가이드",
    "🌱": "마음의 결",
    "📚": "참고",
    "⏰": "시기별",
    "👨‍👩‍👧": "가족과 함께",
    "🩺": "의료 안내",
    "❤️": "사랑의 메시지",
    "🌟": "오늘의 메시지",
    "🌿": "다정한 한 마디",
    "💌": "오늘의 편지",
}


def merge_runs(runs) -> str:
    """Combine adjacent runs with the same bold/italic state.

    docx often splits a single bolded word into multiple runs which produces
    artifacts like "**1**800년". Merging neighbors with the same formatting
    fixes most of these.
    """
    out_parts: list[tuple[str, bool, bool]] = []  # (text, bold, italic)
    for r in runs:
        t = r.text or ""
        if not t:
            continue
        b = bool(r.bold)
        i = bool(r.italic)
        if out_parts and out_parts[-1][1] == b and out_parts[-1][2] == i:
            prev_t, pb, pi = out_parts[-1]
            out_parts[-1] = (prev_t + t, pb, pi)
        else:
            out_parts.append((t, b, i))

    pieces: list[str] = []
    for text, bold, italic in out_parts:
        if not text:
            continue
        # Don't bold/italic pure whitespace
        if not text.strip():
            pieces.append(text)
            continue
        if bold and italic:
            pieces.append(f"***{text}***")
        elif bold:
            pieces.append(f"**{text}**")
        elif italic:
            pieces.append(f"*{text}*")
        else:
            pieces.append(text)
    return "".join(pieces).strip()


def paragraph_to_md(p, level_offset: int = 0) -> str | None:
    """One docx paragraph → one markdown line (or None to drop).

    level_offset: 음수면 헤딩을 한 단계씩 끌어올린다. (예: -1이면 Heading 2가 ##으로,
    Heading 3이 ###으로 매핑되어 챕터 레벨이 ##에 맞춰짐.)
    """
    style = (p.style.name if p.style else "") or ""

    text = merge_runs(p.runs)
    if not text:
        return ""

    # Skip purely decorative ornament-only lines (✨ ✦ 🌱 🌟 등 단독)
    if re.fullmatch(r"[✨✦✧🌱🌟🌿🍃🌸💫⭐·•*\s—–\-]+", text):
        return None

    # 브랜드 정정 — Mothersbaby → 맘곁 (시리즈 사이트의 단일 브랜드 컨텍스트)
    text = re.sub(r"Mothersbaby의?\s*메시지", "맘곁의 메시지", text)
    text = re.sub(r"Mothersbaby", "맘곁", text)

    # 끊어진 bold 정돈 — `**1**800년` → `**1800년**` 형태로 인접 표면적 합치기
    text = re.sub(r"\*\*(\S)\*\*(\S)", r"**\1\2**", text)
    # 인접 bold 결합: **abc**def**ghi** 같은 형태는 그대로 두되,
    # **a****b** → **ab**
    text = re.sub(r"\*\*\s*\*\*", "", text)

    # Heading mapping with optional level shift
    heading_map = {"Heading 1": 2, "Heading 2": 3, "Heading 3": 4, "Heading 4": 5,
                   "Title": 2}
    for prefix, base in heading_map.items():
        if style == prefix or style.startswith(prefix + " "):
            lvl = max(2, base + level_offset)  # never go above ## (H2)
            lvl = min(lvl, 6)
            return f"{'#' * lvl} {text}"

    # 워드 numbering 적용 단락 → bullet
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return f"- {text}"

    # 시각적 불릿 prefix
    if re.match(r"^[•·▪◾◆●▶▷▸▹]\s+", text):
        body = re.sub(r"^[•·▪◾◆●▶▷▸▹]\s+", "", text)
        return f"- {body}"

    # "Chapter 01" / "CHAPTER 1" 단독 라벨 → 드롭
    if re.fullmatch(r"(?i)chapter\s+\d+", text):
        return None
    # 단독 숫자 (목차의 챕터 번호) → 드롭
    if re.fullmatch(r"\d{1,2}", text):
        return None
    # "Ch1. ... Ch2. ... Ch3." 식의 TOC 내비 라인 → 드롭
    if re.search(r"Ch\d+\.", text) and text.count("Ch") >= 2:
        return None
    # "맘곁의 메시지" / "오늘의 메시지" / 비슷한 한 줄 라벨 → 소제목 승격
    if text in ("맘곁의 메시지", "오늘의 메시지", "마음의 메시지"):
        return f"#### 💌 {text}"

    # 이모지 섹션 마커 — 별도 헤더로 승격
    for emoji, default_name in EMOJI_SECTION_MARKERS.items():
        if text.startswith(emoji):
            rest = text[len(emoji):].strip()
            label = rest if rest else default_name
            return f"#### {emoji} {label}"

    return text


def table_to_md(tbl) -> str:
    """docx table → GFM markdown table.

    1셀(1행 1열) 표는 인용 박스로 변환.
    1열 다행 표는 인용 박스(여러 줄)로 변환.
    """
    rows = []
    for row in tbl.rows:
        cells = [c.text.strip().replace("\n", " ").replace("|", "\\|")
                 for c in row.cells]
        rows.append(cells)
    if not rows or not rows[0]:
        return ""

    # 1셀 또는 1열 다행 → blockquote
    if all(len(r) == 1 for r in rows):
        non_empty = [r[0] for r in rows if r[0]]
        if not non_empty:
            return ""
        return "\n".join(f"> {line}" for line in non_empty)

    head = rows[0]
    body = rows[1:]
    out = ["| " + " | ".join(head) + " |"]
    out.append("|" + "|".join(["---"] * len(head)) + "|")
    for r in body:
        if len(r) < len(head):
            r = r + [""] * (len(head) - len(r))
        out.append("| " + " | ".join(r[:len(head)]) + " |")
    return "\n".join(out)


def iter_block_items(parent):
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", child)
        elif child.tag == qn("w:tbl"):
            yield ("tbl", child)


def collapse_split_h2(lines: list[str]) -> list[str]:
    """`## A` 다음 줄이 매우 짧은 단순 텍스트면 합쳐서 H2 한 줄로.

    docx에서 큰 제목이 두 줄로 깨져 들어오는 흔한 패턴을 정돈한다.
    예: `## 태아 이전의` + `마음 준비` → `## 태아 이전의 마음 준비`
    조건: 헤딩 줄 자체도 매우 짧아야 (≤15자). 그래야 부제 단락이 합쳐지는 사고를 막는다.
    """
    out = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("## ") and len(line) <= 18 and i + 1 < len(lines):
            nxt = lines[i + 1]
            if nxt and not nxt.startswith(("#", "- ", "> ", "|")):
                if len(nxt) <= 14:
                    line = line + " " + nxt.strip()
                    out.append(line)
                    i += 2
                    continue
        out.append(line)
        i += 1
    return out


def drop_redundant_meta(lines: list[str], title: str) -> list[str]:
    """첫 부분의 중복 PART/제목 라벨 정리. 동시에 "### 목차" 블록도 드롭."""
    # 1) "### 목차" 또는 "## 목차" 블록 제거 — 다음 챕터(##) 또는 큰 절(### but not 목차 inner)이
    #    나올 때까지 모든 라인 스킵. 본문 빌더 TOC가 자동 생성하므로 중복.
    cleaned = []
    skip_toc_block = False
    for line in lines:
        s = line.strip()
        # 목차 헤더 진입 — 어떤 레벨이든
        if not skip_toc_block and re.fullmatch(r"#{2,4}\s*목차", s):
            skip_toc_block = True
            continue
        if skip_toc_block:
            # 챕터 레벨(##) 헤더가 나오면 토글 해제하고 그 라인은 유지
            if re.match(r"^##\s+(?!목차)", s):
                skip_toc_block = False
                cleaned.append(line)
                continue
            # 그 외(빈 줄, ####, ###, 짧은 항목 등)는 모두 드롭
            continue
        cleaned.append(line)
    lines = cleaned

    out = []
    for line in lines:
        s = line.strip()
        # "새로 쓰는 태교신기" 머리글 제거 (어떤 헤딩 레벨이든)
        bare = re.sub(r"^#+\s*", "", s).replace("**", "").strip()
        if bare.startswith("새로 쓰는 태교신기"):
            continue
        # "태교신기 1.0 | PART N. ..." 같은 헤더 형태 제거
        if re.match(r"^#+\s*태교신기\s*1\.0", s):
            continue
        # "PART N." 으로 시작하는 헤딩 제거 (Title docx 보일러플레이트)
        if re.match(r"^#+\s*PART\s*\d+\.?\s", s):
            continue
        # 단독 "PART N" 또는 "PART N." 라벨 제거
        if re.fullmatch(r"\*?\*?PART\s*\d+\*?\*?\.?", s):
            continue
        # 단독 "프롤로그" / "에필로그" / "여는 글" / "닫는 글" 라벨 제거
        if bare in ("프롤로그", "에필로그", "여는 글", "닫는 글"):
            continue
        out.append(line)
    return out


def drop_duplicate_h2_with_title(lines: list[str], title: str) -> list[str]:
    """첫 # 이후 첫 ## 가 책 제목을 그대로 반복하면 제거."""
    title_norm = re.sub(r"\s+", "", title)
    out = []
    found_h1 = False
    skipped_first_dup_h2 = False
    for i, line in enumerate(lines):
        s = line.strip()
        if not found_h1:
            out.append(line)
            if s.startswith("# "):
                found_h1 = True
            continue
        if not skipped_first_dup_h2 and s.startswith("## "):
            inner = re.sub(r"\s+", "", s[3:])
            if title_norm in inner or inner in title_norm:
                skipped_first_dup_h2 = True
                continue
            skipped_first_dup_h2 = True  # 첫 H2가 다른 거면 유지하고 더 안 본다
        out.append(line)
    return out


def detect_level_offset(doc) -> int:
    """챕터 레벨이 어디인지 자동 감지.

    Heading 1이 여러 개 있으면 chapter = Heading 1 (offset 0).
    Heading 1이 1개뿐이면 chapter = Heading 2 (offset -1).
    """
    h1_count = 0
    for p in doc.paragraphs:
        s = p.style.name if p.style else ""
        if s == "Heading 1" or s.startswith("Heading 1 "):
            h1_count += 1
    return -1 if h1_count <= 1 else 0


def convert(docx_path: Path, out_path: Path, title: str):
    doc = Document(str(docx_path))
    level_offset = detect_level_offset(doc)
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    raw_lines: list[str] = []
    last_blank = True

    def push(s: str):
        nonlocal last_blank
        if s == "":
            if not last_blank:
                raw_lines.append("")
                last_blank = True
            return
        raw_lines.append(s)
        last_blank = False

    push(f"# {title}")
    push("")

    for kind, el in iter_block_items(doc):
        if kind == "p":
            p_obj = para_map.get(el)
            if p_obj is None:
                continue
            md = paragraph_to_md(p_obj, level_offset=level_offset)
            if md is None:
                continue
            push(md)
        elif kind == "tbl":
            t_obj = table_map.get(el)
            if t_obj is None:
                continue
            md = table_to_md(t_obj)
            if md:
                push("")
                push(md)
                push("")

    # 후처리
    raw_lines = drop_redundant_meta(raw_lines, title)
    raw_lines = collapse_split_h2(raw_lines)
    raw_lines = drop_duplicate_h2_with_title(raw_lines, title)

    # Tail blank cleanup
    while raw_lines and not raw_lines[-1].strip():
        raw_lines.pop()

    out_path.write_text("\n".join(raw_lines) + "\n", encoding="utf-8")
    print(f"wrote {out_path}  ({len(raw_lines)} lines)")


def main():
    for docx_name, md_name, title in JOBS:
        src = SRC / docx_name
        if not src.exists():
            print(f"SKIP — not found: {src}", file=sys.stderr)
            continue
        out = OUT / md_name
        convert(src, out, title)


if __name__ == "__main__":
    main()
